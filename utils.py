# Utility Functions for File Conversion

from io import BytesIO
from docx import Document
import fitz  # PyMuPDF
from bs4 import BeautifulSoup
import requests


# Function to convert response to TXT
def convert_to_txt(response):
    return str(response).encode('utf-8', errors='ignore')

# Function to convert response to DOCX
def convert_to_docx(response):
    doc = Document()
    doc.add_heading("Translated", 1)
    doc.add_paragraph(response)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream



# Function to extract text from URL
def extract_text_from_url(url):
    response = requests.get(url)
    response.status_code == 200
    soup = BeautifulSoup(response.text, "html.parser")
    # Extract readable text, here filtering out JavaScript, CSS, and other unwanted tags
    text = ' '.join([p.get_text() for p in soup.find_all("p")])
    return text


# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    text = ""
    with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text += page.get_text("text")
    return text